import pypandoc
import os
import re
import base64
import shutil
from pathlib import Path
from PIL import Image
from docx import Document
from lxml import etree

from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import SystemMessage, HumanMessage
from langchain_core.output_parsers import StrOutputParser
from typing import List

import pymysql
import pymysql.cursors
import chainlit as cl
from google.api_core import exceptions as google_exceptions

# .env 파일이 있다면 로드
#load_dotenv()

# 실제 분석에 사용할 메인 모델
MAIN_MODEL_NAME = "gemini-3-flash-preview"

# API Key 검증용 가벼운 모델 (속도 우선)
#TEST_MODEL_NAME = "gemini-3.1-flash-lite-preview"
TEST_MODEL_NAME = "gemma-4-26b-a4b-it"

async def validate_api_key(api_key: str):
    """
    입력받은 API Key가 유효한지 flash-lite 버전의 LLM으로 빠르게 확인합니다.
    """
    try:
        # 검증용 Lite 모델 생성
        test_llm = ChatGoogleGenerativeAI(model=TEST_MODEL_NAME, google_api_key=api_key, temperature=0)
        # 1토큰 정도의 매우 짧은 응답을 유도하여 연결 확인
        await test_llm.ainvoke([HumanMessage(content="Hi")])
        return True
    except Exception as e:
        print(f"Google API 키가 유효하지 않습니다: {e}")
        return False

# --- 과목별 시스템 프롬프트 정의 (이미지 포함 버전) ---
PROMPTS = {
    "국어": """
<task>국어 시험문제 내용의 꼼꼼하고 정밀한 검토</task>
<specialty>현대 문학, 고전 문학, 문법, 화법과 작문</specialty>
<caution>본문과 선지 내용은 원문의 문구 그대로 검토</caution>
<parsing_status>
    1. 표나 특수 기호는 일부만 parsing
    2. 그림, 사진 등 시각 자료는 텍스트와 함께 제공
    3. 본문에 있는 이미지 파일명(예: 'image1.jpg')을 참고하여 해당 이미지를 함께 검토
</parsing_status>
<to_dos>
    1. 지문 해석, 문법적 규칙, 문학사적 사실 등에서 발생할 수 있는 사실적 오류 여부 점검
    2. 문장을 해석할 때 의미를 오해하거나 중의적으로 해석할 수 있는 여지를 점검
    3. 문제의 질문과 선지가 명확하게 연결되는지, 논리적 비약은 없는지 점검
    4. 선지의 근거를 지문으로부터 명확하게 찾을 수 있는지 점검
    5. 오탈자나 어색한 문장 구조 점검
    6. 객관식 문제라면 선지 번호 가운데 문제에 대한 정답 존재 여부 점검
    7. 선다형 문제라면 선지 번호 중 ㄱ, ㄴ, ㄷ을 조합하여 문제에 대한 정답 존재 여부 점검
    8. 서술형, 논술형 문항을 포함한 모든 문항에 대하여 검토 의견 제시
</to_dos>
<output_instructions>
    1. 반드시 아래의 마크다운 형식을 지켜서 답변.
    2. 서론이나 잡담 없이 바로 본론 작성.
</output_instructions>
<output_example>
## [문항별 검토]
### 문항 1
- **검토 결과**: [적절 / 수정 필요 / 오류]
- **상세 내용**: (To Dos의 모든 지시사항과 관련된 내용 작성)

### 문항 2
...

## [주요 수정 사항 요약]
- (반드시 고쳐야 할 부분 리스트)
</output_example>""",
    "수학": """
<task>수학 시험문제의 꼼꼼하고 정밀한 검토</task>
<specialty>미적분, 확률과 통계, 기하, 대수</specialty>
<caution>본문과 선지 내용은 원문의 문구 그대로 검토</caution>
<parsing_status>
    1. 표나 특수 기호는 일부만 parsing
    2. 그림, 사진 등 시각 자료는 텍스트와 함께 제공
    3. 본문에 있는 이미지 파일명(예: 'image1.jpg')을 참고하여 해당 이미지를 함께 검토
</parsing_status>
<to_dos>
    1. 문제에 사용된 수학적 개념, 공식, 정의의 오류 여부 점검
    2. 문제 풀이 과정에서 발생할 수 있는 논리적 오류나 계산 실수를 검토 (단, 직접 풀이를 제공하지는 않음)
    3. 질문의 조건이 명확하고 충분한지 점검
    4. 해가 유일하게 결정되는지 점검
    5. 용어나 기호를 중의적으로 해석할 수 있는 경우를 점검
    6. 오탈자나 어색한 문장 구조 점검
    6. 객관식 문제라면 선지 번호 가운데 문제에 대한 정답 존재 여부 점검
    7. 선다형 문제라면 선지 번호 중 ㄱ, ㄴ, ㄷ을 조합하여 문제에 대한 정답 존재 여부 점검
    8. 서술형, 논술형 문항을 포함한 모든 문항에 대하여 검토 의견 제시
</to_dos>
<output_instructions>
    1. 반드시 아래의 마크다운 형식을 지켜서 답변.
    2. 서론이나 잡담 없이 바로 본론 작성.
</output_instructions>
<output_example>
## [문항별 검토]
### 문항 1
- **검토 결과**: [적절 / 수정 필요 / 오류]
- **상세 내용**: (To Dos의 모든 지시사항과 관련된 내용 작성)

### 문항 2
...

## [주요 수정 사항 요약]
- (반드시 고쳐야 할 부분 리스트)
</output_example>""",
    "영어": """
<task>영어 시험문제의 꼼꼼하고 정밀한 검토</task>
<specialty>영문법, 독해, 어휘, 작문</specialty>
<caution>본문과 선지 내용은 원문의 문구 그대로 검토</caution>
<parsing_status>
    1. 표나 특수 기호는 일부만 parsing
    2. 그림, 사진 등 시각 자료는 텍스트와 함께 제공
    3. 본문에 있는 이미지 파일명(예: 'image1.jpg')을 참고하여 해당 이미지를 함께 검토
</parsing_status>
<to_dos>
    1. 문법적 오류(Grammatical errors)가 있는지 점검
    2. 어휘 선택이 부적절하거나(Inappropriate vocabulary), 문맥에 맞지 않는 단어가 사용되었는지 점검
    3. 문장을 해석할 때 의미를 오해하거나 중의적으로 해석할 수 있는 애매한(ambiguous) 표현을 점검
    4. 지문의 내용과 질문, 선지 사이의 논리적 일관성을 점검
    5. 오탈자(typo)나 구두점 오류 점검
    6. 객관식 문제라면 선지 번호 가운데 문제에 대한 정답 존재 여부 점검
    7. 선다형 문제라면 선지 번호 중 ㄱ, ㄴ, ㄷ을 조합하여 문제에 대한 정답 존재 여부 점검
    8. 서술형, 논술형 문항을 포함한 모든 문항에 대하여 검토 의견 제시
</to_dos>
<output_instructions>
    1. 반드시 아래의 마크다운 형식을 지켜서 답변.
    2. 서론이나 잡담 없이 바로 본론 작성.
</output_instructions>
<output_example>
## [문항별 검토]
### 문항 1
- **검토 결과**: [적절 / 수정 필요 / 오류]
- **상세 내용**: (To Dos의 모든 지시사항과 관련된 내용 작성)

### 문항 2
...

## [주요 수정 사항 요약]
- (반드시 고쳐야 할 부분 리스트)
</output_example>""",
    "사회": """
<task>사회·윤리·역사 시험문제의 꼼꼼하고 정밀한 검토</task>
<specialty>사회·문화, 경제, 윤리, 철학, 한국사, 세계사, 정치와 법</specialty>
<caution>본문과 선지 내용은 원문의 문구 그대로 검토</caution>
<parsing_status>
    1. 표나 특수 기호는 일부만 parsing
    2. 그림, 사진 등 시각 자료는 텍스트와 함께 제공
    3. 본문에 있는 이미지 파일명(예: 'image1.jpg')을 참고하여 해당 이미지를 함께 검토
</parsing_status>
<to_dos>
    1. 법률 조항, 사회 과학적 개념, 윤리, 철학, 역사적 사실 등 내용상의 사실 오류(factual errors)를 점검
    2. 특정 관점에 편향되었거나 논쟁의 소지가 큰 표현이 있는지 점검
    3. 문장을 해석할 때 의미를 오해하거나 중의적으로 해석할 수 있는 여지를 점검
    4. 연도, 인물, 사건 등의 명칭에 오탈자가 있는지 점검
    5. 오탈자(typo)나 구두점 오류 점검
    6. 객관식 문제라면 선지 번호 가운데 문제에 대한 정답 존재 여부 점검
    7. 선다형 문제라면 선지 번호 중 ㄱ, ㄴ, ㄷ을 조합하여 문제에 대한 정답 존재 여부 점검
    8. 서술형, 논술형 문항을 포함한 모든 문항에 대하여 검토 의견 제시
</to_dos>
<output_instructions>
    1. 반드시 아래의 마크다운 형식을 지켜서 답변.
    2. 서론이나 잡담 없이 바로 본론 작성.
</output_instructions>
<output_example>
## [문항별 검토]
### 문항 1
- **검토 결과**: [적절 / 수정 필요 / 오류]
- **상세 내용**: (To Dos의 모든 지시사항과 관련된 내용 작성)

### 문항 2
...

## [주요 수정 사항 요약]
- (반드시 고쳐야 할 부분 리스트)
</output_example>""",
    "과학": """
<task>과학 시험문제의 꼼꼼하고 정밀한 검토</task>
<specialty>물리학, 화학, 생명과학, 지구과학</specialty>
<caution>본문과 선지 내용은 원문의 문구 그대로 검토</caution>
<parsing_status>
    1. 표나 특수 기호는 일부만 parsing
    2. 그림, 사진 등 시각 자료는 텍스트와 함께 제공
    3. 본문에 있는 이미지 파일명(예: 'image1.jpg')을 참고하여 해당 이미지를 함께 검토
</parsing_status>
<to_dos>
    1. 문제에 사용된 과학적 개념, 공식, 정의의 오류 여부 점검
    2. 문장의 의미 속에 내재된 과학적인 오류 여부를 점검
    3. 문장 자체에 내재된 언어학적인 오류 여부를 점검
    4. 과학 용어나 문장을 중의적으로 해석할 수 있는 경우를 점검
    5. 오탈자(typo)나 구두점 오류 점검
    6. 객관식 문제라면 선지 번호 가운데 문제에 대한 정답 존재 여부 점검
    7. 선다형 문제라면 선지 번호 중 ㄱ, ㄴ, ㄷ을 조합하여 문제에 대한 정답 존재 여부 점검
    8. 서술형, 논술형 문항을 포함한 모든 문항에 대하여 검토 의견 제시
</to_dos>
<output_instructions>
    1. 반드시 아래의 마크다운 형식을 지켜서 답변.
    2. 서론이나 잡담 없이 바로 본론 작성.
</output_instructions>
<output_example>
## [문항별 검토]
### 문항 1
- **검토 결과**: [적절 / 수정 필요 / 오류]
- **상세 내용**: (To Dos의 모든 지시사항과 관련된 내용 작성)
- **정답 존재 여부**: [이상 없음 / 확인 필요]

### 문항 2
...

## [주요 수정 사항 요약]
- (반드시 고쳐야 할 부분 리스트)
</output_example>""",
}

# --- DB 처리 함수 ---

# --- 수정된 DB 처리 함수 (과목 컬럼 분리 버전) ---

def log_usage_history(session_id, subject, is_success, error_message=None):
    """
    사용 이력을 저장합니다.
    - exam_content: 보안상 "CONTENT_NOT_STORED" 저장
    - llm_review: 성공/실패 상태 메시지 저장
    - subject: 별도 컬럼에 과목명 저장
    """
    connection = None
    try:
        connection = pymysql.connect(
            host=os.getenv("MARIADB_HOST"),
            user=os.getenv("MARIADB_USER"),
            password=os.getenv("MARIADB_PASSWORD"),
            db=os.getenv("MARIADB_DATABASE"),
            charset='utf8mb4',
            cursorclass=pymysql.cursors.DictCursor,
            port=int(os.getenv("MARIADB_PORT"))
        )
        with connection.cursor() as cursor:
            # 보안 처리
            safe_content_log = "CONTENT_NOT_STORED (Security Policy)"
            
            # 상태 메시지 (과목명은 별도 컬럼에 저장하므로 여기선 상태만 기록해도 됨)
            if is_success:
                status_log = "[SUCCESS] Review Completed"
            else:
                status_log = f"[ERROR] {error_message}"

            # 쿼리 수정: subject 컬럼 추가
            sql = """
                INSERT INTO paper_review_logs (session_id, subject, exam_content, llm_review) 
                VALUES (%s, %s, %s, %s)
            """
            cursor.execute(sql, (session_id, subject, safe_content_log, status_log))
            
        connection.commit()
        print(f"이력 저장 완료: session_id={session_id}, subject={subject}, success={is_success}")
    except pymysql.Error as e:
        print(f"이력 저장 실패: {e}")
    finally:
        if connection:
            connection.close()

# --- 파일 및 이미지 처리 함수 ---

def debug_math_structure(oMath_element):
    """수식의 XML 구조를 출력 (디버깅용)"""
    print("\n    [수식 XML 구조]:")
    xml_str = etree.tostring(oMath_element, encoding='unicode', pretty_print=True)
    # 처음 500자만 출력
    print(xml_str[:500])
    print("    ...")


def omml_to_text(oMath_element, debug=False):
    """OMML 수식을 텍스트로 변환"""
    M_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
    
    if debug:
        debug_math_structure(oMath_element)
    
    # 방법 1: m:t 요소에서 텍스트 추출
    text_parts = []
    for t in oMath_element.iter(f'{M_NS}t'):
        if t.text:
            text_parts.append(t.text)
            if debug:
                print(f"      m:t 발견: '{t.text}'")
    
    if text_parts:
        result = ''.join(text_parts)
        if debug:
            print(f"      결과 (m:t): '{result}'")
        return result
    
    # 방법 2: 모든 텍스트 노드 추출
    all_text = ''.join(oMath_element.itertext()).strip()
    if all_text:
        if debug:
            print(f"      결과 (itertext): '{all_text}'")
        return all_text
    
    if debug:
        print("      결과: (빈 문자열)")
    
    return ''

# ============================================================
# 1. OMML(수식)을 LaTeX로 변환하는 함수 정의 (먼저 실행되어야 함)
# ============================================================
def omml_to_latex(element):
    """
    OMML(Office Math Markup Language) 요소를 LaTeX 문자열로 변환하는 재귀 함수
    """
    namespaces = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    
    tag = element.tag
    local_tag = tag.split('}')[-1] # 네임스페이스 제거한 태그명

    # 1. 텍스트 (m:t)
    if local_tag == 't':
        return element.text if element.text else ""

    # 2. 분수 (m:f) -> \frac{분자}{분모}
    elif local_tag == 'f':
        num_node = element.find('m:num', namespaces)
        den_node = element.find('m:den', namespaces)
        num = omml_to_latex(num_node) if num_node is not None else ""
        den = omml_to_latex(den_node) if den_node is not None else ""
        return f"\\frac{{{num}}}{{{den}}}"

    # 3. 위첨자 (m:sSup) -> ^{...}
    elif local_tag == 'sSup':
        e_node = element.find('m:e', namespaces)
        sup_node = element.find('m:sup', namespaces)
        base = omml_to_latex(e_node) if e_node is not None else ""
        sup = omml_to_latex(sup_node) if sup_node is not None else ""
        return f"{{{base}}}^{{{sup}}}"

    # 4. 아래첨자 (m:sSub) -> _{...}
    elif local_tag == 'sSub':
        e_node = element.find('m:e', namespaces)
        sub_node = element.find('m:sub', namespaces)
        base = omml_to_latex(e_node) if e_node is not None else ""
        sub = omml_to_latex(sub_node) if sub_node is not None else ""
        return f"{{{base}}}_{{{sub}}}"

    # 5. 위아래첨자 (m:sSubSup) -> _{...}^{...}
    elif local_tag == 'sSubSup':
        e_node = element.find('m:e', namespaces)
        sub_node = element.find('m:sub', namespaces)
        sup_node = element.find('m:sup', namespaces)
        base = omml_to_latex(e_node) if e_node is not None else ""
        sub = omml_to_latex(sub_node) if sub_node is not None else ""
        sup = omml_to_latex(sup_node) if sup_node is not None else ""
        return f"{{{base}}}_{{{sub}}}^{{{sup}}}"

    # 6. 근호/루트 (m:rad) -> \sqrt[...]
    elif local_tag == 'rad':
        deg_node = element.find('m:deg', namespaces) # 제곱근의 차수
        e_node = element.find('m:e', namespaces)
        base = omml_to_latex(e_node) if e_node is not None else ""
        
        if deg_node is not None and len(deg_node) > 0:
            deg = omml_to_latex(deg_node)
            if deg.strip():
                return f"\\sqrt[{deg}]{{{base}}}"
        return f"\\sqrt{{{base}}}"
    
    # 7. 괄호/구분자 (m:d) -> \left( ... \right)
    elif local_tag == 'd':
        dPr = element.find('m:dPr', namespaces)
        begChr = '('
        endChr = ')'
        if dPr is not None:
            beg_attr = dPr.find('m:begChr', namespaces)
            end_attr = dPr.find('m:endChr', namespaces)
            if beg_attr is not None: begChr = beg_attr.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val') or '('
            if end_attr is not None: endChr = end_attr.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val') or ')'
            
        e_nodes = element.findall('m:e', namespaces)
        content = "".join([omml_to_latex(child) for child in e_nodes])
        return f"\\left{begChr} {content} \\right{endChr}"

    # 기본 재귀 호출
    else:
        text = ""
        for child in element:
            text += omml_to_latex(child)
        return text

# ============================================================
# 2. 단락 내용 추출 함수 (omml_to_latex를 호출함)
# ============================================================
def extract_paragraph_content(para_element, debug=False):
    """단락에서 텍스트와 수식을 순서대로 추출"""
    
    # 네임스페이스 정의
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    }
    
    if debug:
        print("\n  [단락 처리 시작]")
    
    para_parts = []
    
    # 단락의 모든 자식 요소 순회
    for child in para_element:
        # 수식(OMML) 발견 시
        if child.tag.endswith('oMath') or child.tag.endswith('oMathPara'):
            # 여기서 위에서 정의한 omml_to_latex 함수를 호출합니다
            latex_eq = omml_to_latex(child)
            para_parts.append(f" ${latex_eq}$ ") 
        
        # 일반 텍스트 런(w:r) 처리
        elif child.tag.endswith('r'):
            t_vals = child.findall('.//w:t', namespaces)
            for t in t_vals:
                if t.text:
                    para_parts.append(t.text)
    
    result = ''.join(para_parts).strip()
    
    if debug:
        print(f"  [단락 결과]: '{result[:100]}{'...' if len(result) > 100 else ''}'")
    
    return result


def extract_table_from_element(table_element, debug=False):
    """XML 테이블 요소에서 마크다운 표 생성"""
    W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    if debug:
        print("\n[표 처리 시작]")
    
    rows = []
    for row_idx, tr in enumerate(table_element.iter(f'{W_NS}tr')):
        row = []
        for col_idx, tc in enumerate(tr.iter(f'{W_NS}tc')):
            cell_parts = []
            
            # 셀 내의 각 단락 처리
            for para in tc.findall(f'{W_NS}p'):
                if debug:
                    print(f"  표 셀 [{row_idx},{col_idx}]:")
                para_text = extract_paragraph_content(para, debug=debug)
                if para_text:
                    cell_parts.append(para_text)
            
            cell_text = ' '.join(cell_parts).strip()
            row.append(cell_text if cell_text else ' ')
        
        if row:
            rows.append(row)
    
    if not rows:
        return None
    
    # 마크다운 표 생성
    markdown_table = []
    
    if rows:
        header = ' | '.join(rows[0])
        markdown_table.append(f"| {header} |")
        
        separator = ' | '.join(['---'] * len(rows[0]))
        markdown_table.append(f"| {separator} |")
        
        for row in rows[1:]:
            while len(row) < len(rows[0]):
                row.append(' ')
            row_text = ' | '.join(row[:len(rows[0])])
            markdown_table.append(f"| {row_text} |")
    
    if debug:
        print("[표 처리 완료]")
    
    return '\n'.join(markdown_table)


def extract_textbox_content(textbox_element, debug=False):
    """글상자 요소에서 텍스트, 표, 수식을 모두 추출"""
    W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    content_parts = []
    
    if debug:
        print("\n" + "="*60)
        print("글상자 내용 추출 시작")
        print("="*60)
    
    # w:txbxContent 찾기
    for txbxContent in textbox_element.iter(f'{W_NS}txbxContent'):
        for child in txbxContent:
            # 표 발견
            if child.tag == f'{W_NS}tbl':
                table_md = extract_table_from_element(child, debug=debug)
                if table_md:
                    content_parts.append('\n' + table_md + '\n')
            
            # 단락 발견
            elif child.tag == f'{W_NS}p':
                para_text = extract_paragraph_content(child, debug=False)
                if para_text:
                    content_parts.append(para_text)
    
    result = '\n\n'.join(content_parts)
    
    if debug:
        print("="*60)
        print(f"글상자 내용 추출 완료 ({len(result)}자)")
        print("="*60 + "\n")
    
    return result


def convert_docx_to_text_with_images(file_path, media_dir='./extracted_media', debug=False):
    """글상자를 포함한 DOCX를 마크다운으로 변환"""
    
    if not os.path.exists(file_path):
        return "파일을 찾을 수 없습니다."
    
    try:
        print(f"변환 시작: {file_path}")
        
        doc = Document(file_path)
        textbox_map = {}
        all_textboxes = []
        
        temp_path = file_path.replace('.docx', '_temp.docx')
        
        # 네임스페이스 정의
        VML_NS = '{urn:schemas-microsoft-com:vml}'
        WPS_NS = '{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}'
        
        textbox_count = 0
        for para in doc.paragraphs:
            para_textboxes = []
            para_element = para._element
            
            # VML textbox
            for textbox in para_element.iter(f'{VML_NS}textbox'):
                textbox_count += 1
                if debug:
                    print(f"\n>>> 글상자 #{textbox_count} (VML) 처리 중...")
                
                content = extract_textbox_content(textbox, debug=debug)
                
                if content.strip():
                    para_textboxes.append(content)
                    all_textboxes.append(content)
            
            # wps:txbx
            for txbx in para_element.iter(f'{WPS_NS}txbx'):
                textbox_count += 1
                if debug:
                    print(f"\n>>> 글상자 #{textbox_count} (WPS) 처리 중...")
                
                content = extract_textbox_content(txbx, debug=debug)
                
                if content.strip():
                    para_textboxes.append(content)
                    all_textboxes.append(content)
            
            if para_textboxes:
                textbox_map[id(para)] = para_textboxes
                for content in para_textboxes:
                    preview = content[:80].replace('\n', ' ') + "..." if len(content) > 80 else content.replace('\n', ' ')
                    if not debug:  # debug 모드가 아닐 때만 간단한 요약 출력
                        print(f"  ✓ 글상자 발견: {preview}")
        
        print(f"\n총 {len(all_textboxes)}개의 글상자를 발견했습니다.")
        
        if not all_textboxes:
            print("글상자를 찾지 못했습니다. Pandoc만으로 변환합니다.")
            markdown = pypandoc.convert_file(
                file_path,
                to='markdown',
                format='docx',
                extra_args=[
                    '--wrap=none',
                    '--standalone',
                    f'--extract-media={media_dir}'
                ]
            )
            return markdown
        
        # 마커 삽입
        marker_to_content = {}
        inserted_count = 0
        
        for para in doc.paragraphs:
            para_id = id(para)
            if para_id in textbox_map:
                textboxes = textbox_map[para_id]
                
                for content in textboxes:
                    marker = f"@@TEXTBOX{inserted_count + 1}@@"
                    marker_to_content[marker] = content
                    
                    if para.runs:
                        para.runs[-1].text += f" {marker} "
                    else:
                        para.add_run(f" {marker} ")
                    
                    inserted_count += 1
                    if not debug:
                        preview = content[:50].replace('\n', ' ') + "..." if len(content) > 50 else content.replace('\n', ' ')
                        print(f"  ✓ {marker} 삽입됨: {preview}")
        
        print(f"\n마커 삽입 완료: {inserted_count}개")
        
        doc.save(temp_path)
        
        # Pandoc 변환
        markdown = pypandoc.convert_file(
            temp_path,
            to='markdown',
            format='docx',
            extra_args=[
                '--wrap=none',
                '--standalone',
                f'--extract-media={media_dir}'
            ]
        )
        
        # 마커 확인
        found = sum(1 for m in marker_to_content if m in markdown)
        print(f"  발견된 마커: {found}/{len(marker_to_content)}개")
        
        # 마커 치환
        replaced_count = 0
        
        for marker, content in marker_to_content.items():
            if marker in markdown:
                if '|' in content and '---' in content:
                    formatted_content = f"\n\n> 📦 **[글상자]**\n\n{content}\n\n"
                else:
                    lines = content.split('\n')
                    formatted_lines = ['> ' + line if line.strip() else '>' for line in lines]
                    formatted_content = f"\n\n> 📦 **[글상자]**\n" + '\n'.join(formatted_lines) + "\n\n"
                
                markdown = markdown.replace(marker, formatted_content)
                replaced_count += 1
                if not debug:
                    print(f"  ✓ {marker} 치환 완료")
        
        # 미치환 글상자 처리
        if replaced_count < len(marker_to_content):
            print(f"\n⚠ {len(marker_to_content) - replaced_count}개 글상자를 문서 끝에 추가")
            markdown += "\n\n---\n\n## 📦 글상자 내용 (원래 위치 확인 불가)\n\n"
            
            for i, (marker, content) in enumerate(marker_to_content.items(), 1):
                if marker not in markdown:
                    if '|' in content and '---' in content:
                        markdown += f"\n### 글상자 {i}\n\n{content}\n\n"
                    else:
                        lines = content.split('\n')
                        formatted_lines = ['> ' + line if line.strip() else '>' for line in lines]
                        markdown += f"\n### 글상자 {i}\n\n" + '\n'.join(formatted_lines) + "\n\n"
        
        if os.path.exists(temp_path):
            os.remove(temp_path)
        
        print(f"\n✅ 변환 완료: {replaced_count}/{len(marker_to_content)}개 글상자 삽입됨")
        
        return markdown
        
    except Exception as e:
        temp_path = file_path.replace('.docx', '_temp.docx')
        if os.path.exists(temp_path):
            os.remove(temp_path)
        import traceback
        return f"오류 발생: {e}\n\n{traceback.format_exc()}"

def preprocess_content(text):
    """
    Pandoc으로 변환된 텍스트의 노이즈를 제거하고 이미지 태그를 정규화합니다.
    1. 불필요한 ASCII 테이블 선(+---+) 제거
    2. 복잡한 이미지 경로 및 속성 제거 -> [이미지: image1.jpg] 형태로 통일
    3. 과도한 공백 제거
    """

    # 1. ASCII 테이블의 구분선 (+----+) 제거
    # 줄의 시작이 +이고, 그 뒤로 -나 +로만 이루어진 줄을 찾아 제거
    text = re.sub(r'^\+[-+]+$', '', text, flags=re.MULTILINE)

    # 2. 이미지 태그 정규화 및 확장자 강제 변경 (.jpg)
    # 패턴: ![설명](경로){속성} 또는 ![설명](경로)
    def clean_image_tag(match):
        # 정규식 그룹 1번: 이미지 경로 (예: ./extracted_media/media/image1.png)
        full_path = match.group(1)
        
        # 파일명만 추출 (image1.png)
        filename = os.path.basename(full_path)
        name, ext = os.path.splitext(filename)
        
        # 사용자의 로직에 맞춰 모든 확장자를 .jpg로 변경하여 표기
        new_filename = f"{name}.jpg"
        
        # LLM이 인식하기 쉬운 형태로 변경
        return f"\n[참고 이미지: {new_filename}]\n"

    # Markdown 이미지 링크 패턴 매칭 및 치환
    # ![](path){opt} 또는 ![](path) 형태 모두 대응
    text = re.sub(r'!\[.*?\]\((.*?)\)(?:\{.*?\})?', clean_image_tag, text)

    # 3. 테이블 파이프(|)로 인한 불필요한 공백 및 줄바꿈 정리
    # 연속된 3개 이상의 줄바꿈을 2개로 축소
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text

def resize_image_if_needed(img, max_width=640):
    """이미지 가로 크기가 max_width보다 크면 비율 유지하며 리사이징"""
    width, height = img.size
    if width > max_width:
        ratio = max_width / width
        new_height = int(height * ratio)
        img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
    return img

def convert_to_grayscale(img):
    """이미지를 그레이스케일로 변환"""
    if img.mode != 'L':
        img = img.convert('L')
    return img

def convert_images_to_jpg(directory, max_width=640):
    """EMF, PNG, BMP, WMF 파일을 JPG로 변환 (리사이징 및 그레이스케일)"""
    directory = Path(directory)
    if not directory.exists(): return
    extensions = ['.emf', '.png', '.bmp', '.wmf']
    for ext in extensions:
        for file_path in directory.glob(f'*{ext}'):
            try:
                with Image.open(file_path) as img:
                    if img.mode in ('RGBA', 'LA', 'P'):
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                        img = background
                    elif img.mode not in ('RGB', 'L'):
                        img = img.convert('RGB')
                    
                    img = resize_image_if_needed(img, max_width)
                    img = convert_to_grayscale(img)
                    
                    new_file_path = file_path.with_suffix('.jpg')
                    img.save(new_file_path, 'JPEG', quality=75, optimize=True)
            except Exception as e:
                print(f"이미지 변환 실패: {file_path.name} - {str(e)}")

def resize_existing_jpg(directory, max_width=640):
    """기존 JPG/JPEG 파일을 리사이징 및 그레이스케일 변환"""
    directory = Path(directory)
    if not directory.exists(): return
    jpg_files = list(directory.glob('*.jpg')) + list(directory.glob('*.jpeg'))
    for file_path in jpg_files:
        try:
            with Image.open(file_path) as img:
                needs_processing = img.size[0] > max_width or img.mode != 'L'
                if needs_processing:
                    img = resize_image_if_needed(img, max_width)
                    img = convert_to_grayscale(img)
                    img.save(file_path, 'JPEG', quality=75, optimize=True)
        except Exception as e:
            print(f"JPG 처리 실패: {file_path.name} - {str(e)}")

def encode_images_to_base64(directory):
    """디렉토리의 모든 JPG/JPEG 파일을 Base64로 인코딩 (파일명 포함)"""
    directory = Path(directory)
    if not directory.exists(): return []
    encoded_data = []
    jpg_files = sorted(list(directory.glob('*.jpg'))) + sorted(list(directory.glob('*.jpeg')))
    for file_path in jpg_files:
        try:
            with open(file_path, 'rb') as image_file:
                encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
                encoded_data.append({
                    "filename": file_path.name,  # ← 파일명 추가
                    "b64": encoded_string
                })
        except Exception as e:
            print(f"Base64 인코딩 실패: {file_path.name} - {str(e)}")
    return encoded_data

# --- Chainlit 앱 로직 ---

@cl.on_chat_start
async def on_chat_start():
    res = await cl.AskUserMessage(
        content="안녕하세요? 학교 시험문제 검토 도우미입니다.\n시험문제 내용과 검토 결과 모두 서버에 전혀 저장하지 않도록 설계했으니 걱정 말고 마음껏 검토하세요.\n\n"
                "**Google Gemini API 키**를 입력해주세요.\n\n"
                "API 키도 서버에 전혀 기록하지 않습니다. 안심하고 입력하세요..\n"
                "무료 API 키 하나당 **하루 최대 20회**까지 시험 문제 파일을 통째로 올려서 검토할 수 있습니다.\n"
                "Google Gemini API 키는 https://aistudio.google.com/app/api-keys 에서 무료로 발급받을 수 있습니다.", 
        timeout=600
    ).send()

    if res:
        user_api_key = res["output"].strip()
        
        # 검증 중 메시지
        msg = cl.Message(content=f"**{TEST_MODEL_NAME}** 모델로 API 키 연결 확인 중...", author="문제 검토 도우미")
        await msg.send()
        
        # Lite 모델로 빠르게 검증
        is_valid = await validate_api_key(user_api_key)
        
        # 수정안: 재입력 루프 없이 검증 실패 시 안내만 하고 종료
        if is_valid:
            cl.user_session.set("user_api_key", user_api_key)
            msg.content = "✅ API Key가 확인되었습니다."
            await msg.update()
        else:
            msg.content = "❌ 유효하지 않은 API Key입니다. 페이지를 새로고침 후 다시 시도해주세요."
            await msg.update()
            return
    else:
        # 타임아웃
        await cl.Message(content="입력 시간이 초과되었습니다. 페이지를 새로고침 해주세요.").send()
        return

    # 2. 인사말 및 영역 선택
    await cl.Message(content=f"분석 모델: {MAIN_MODEL_NAME}", author="문제 검토 도우미").send()

    actions = [
        cl.Action(name="subject_select", value=subject, label=subject, payload={"subject": subject})
        for subject in PROMPTS.keys()
    ]
    await cl.Message(
        content="검토할 과목을 선택해주세요.",
        actions=actions,
        author="문제 검토 도우미"
    ).send()


@cl.action_callback("subject_select")
async def on_subject_select(action: cl.Action):
    """과목 선택 시 세션에 저장하고 안내를 표시합니다."""
    subject = action.payload["subject"]
    cl.user_session.set("subject", subject)
    cl.user_session.set("session_id", cl.user_session.get("id"))
    await cl.Message(content=f"**{subject}** 과목을 선택했습니다.").send()
    await cl.Message(content="사용법을 안내할게요.\n1. 한/글에서 시험문제 hwp 파일을 여신 후, **암호를 해제**합니다.\n2. 파일 - 다른 이름으로 저장 - 파일 형식으로 **워드 문서 (*.docx)** 로 저장합니다.\n3. **변환한 시험문제 파일(.docx)을 열어서 불필요하게 생성된 문제 번호를 지우고 저장**합니다. 한/글에서 워드 문서로의 변환이 완벽히 되지 않기 때문입니다.\n4. 변환한 시험문제 파일(.docx)을 업로드해주세요.", author="문제 검토 도우미").send()
    await action.remove()


@cl.on_message
async def on_message(message: cl.Message):
    # on_message 함수 내부에서
    user_api_key = cl.user_session.get("user_api_key")
    if not user_api_key:
        await cl.Message(content="API 키가 없습니다. 페이지를 새로고침 해주세요.").send()
        return

    llm = ChatGoogleGenerativeAI(
        model=MAIN_MODEL_NAME,
        google_api_key=user_api_key,
        thinking_level="high"
    )

    subject = cl.user_session.get("subject")
    if not subject:
        await cl.Message(content="출제한 시험문제의 과목을 선택해주세요.", author="문제 검토 도우미").send()
        return

    files = [f for f in message.elements if isinstance(f, cl.File)]
    if not files:
        await cl.Message(content="변환한 docx 파일을 업로드해주세요.", author="문제 검토 도우미").send()
        return

    uploaded_file = files[0]
    if not uploaded_file.name.lower().endswith(".docx"):
        await cl.Message(content="docx 파일만 업로드할 수 있습니다.", author="문제 검토 도우미").send()
        return

    session_id = cl.user_session.get("session_id")
    media_dir = f"./extracted_media_{session_id}"
    
    if os.path.exists(media_dir):
        shutil.rmtree(media_dir)
    os.makedirs(media_dir)

    try:
        await cl.Message(content=f"'{uploaded_file.name}' 파일의 검토를 시작합니다.\n\n텍스트와 이미지를 함께 처리합니다. 단, LLM의 특성상 이미지 처리가 완벽하지 않고, 글상자 속에 내용이 들어 있다면 불완전하게 추출된 채로 자료를 처리하는 기술적 한계를 감안하여 검토 의견을 읽어 주세요.\n\n답변이 나오기까지는 최대 3분 정도 걸립니다.", author="문제 검토 도우미").send()
        
        extracted_text = await cl.make_async(convert_docx_to_text_with_images)(uploaded_file.path, media_dir, debug=False)
        
        if extracted_text.startswith("오류:"):
            # 변환 실패 로그 저장
            await cl.make_async(log_usage_history)(session_id, subject, False, error_message=f"File Convert Error: {extracted_text}")
            await cl.Message(content=f"파일 처리 중 문제가 발생했습니다.\n\n{extracted_text}", author="문제 검토 도우미").send()
            return
        
        extracted_text = preprocess_content(extracted_text)
        image_processing_dir = os.path.join(media_dir, 'media')

        await cl.make_async(convert_images_to_jpg)(image_processing_dir)
        await cl.make_async(resize_existing_jpg)(image_processing_dir)
        encoded_images = await cl.make_async(encode_images_to_base64)(image_processing_dir)

        system_role = PROMPTS[subject]
        
        human_contents = [{"type": "text", "text": f"시험지 본문 내용: {extracted_text}"}]

        for item in encoded_images:
            # 어떤 이미지인지 LLM에게 명확히 알려줌
            human_contents.append({
                "type": "text",
                "text": f"[이미지 파일명: {item['filename']}]"
            })
            human_contents.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{item['b64']}"}
            })
        
        messages = [
            SystemMessage(content=system_role),
            HumanMessage(content=human_contents)
        ]

        msg = cl.Message(content="", author="검토 의견")
        await msg.send()

        full_response = []
        async with cl.Step(name="시험문제 검토 중..."):
            async for chunk in llm.astream(messages):
                if isinstance(chunk.content, list):
                    # 리스트인 경우 텍스트만 추출
                    content = ''.join([
                        item.get('text', '') if isinstance(item, dict) else str(item)
                        for item in chunk.content
                    ])
                elif isinstance(chunk.content, str):
                    content = chunk.content
                else:
                    # 예상치 못한 타입은 문자열로 변환
                    content = str(chunk.content)
                
                if content:  # 빈 문자열이 아닐 때만 처리
                    await msg.stream_token(content)
                    full_response.append(content)
            
            await msg.update()

        # [수정됨] 정상 완료 시 로그 저장 (성공=True, 에러메시지=None)
        await cl.make_async(log_usage_history)(session_id, subject, True)

    except google_exceptions.InvalidArgument as e:
        error_msg = str(e)
        if "request is too large" in error_msg or "token" in error_msg.lower():
            user_msg = "오류: 입력 내용이 너무 깁니다. API가 처리할 수 있는 최대 토큰 길이를 초과했습니다. DOCX 파일을 여러 개로 나누어 업로드해주세요."
            await cl.Message(content=user_msg, author="오류").send()
        else:
            await cl.Message(content=f"API 처리 중 오류가 발생했습니다: {e}", author="오류").send()
        
        # [수정됨] API 오류 시 로그 저장 (성공=False, 에러메시지 기록)
        await cl.make_async(log_usage_history)(session_id, subject, False, error_message=f"Google API Error: {error_msg}")

    except Exception as e:
        error_msg = str(e)
        await cl.Message(content=f"처리 중 예기치 않은 오류가 발생했습니다: {e}", author="문제 검토 도우미").send()
        
        # [수정됨] 일반 오류 시 로그 저장 (성공=False, 에러메시지 기록)
        await cl.make_async(log_usage_history)(session_id, subject, False, error_message=f"General Exception: {error_msg}")

    finally:
        if os.path.exists(media_dir):
            shutil.rmtree(media_dir)
            print(f"임시 디렉토리 삭제 완료: {media_dir}")

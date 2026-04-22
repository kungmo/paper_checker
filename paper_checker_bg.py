import pypandoc
import os
import re
import base64
import shutil
from pathlib import Path
from PIL import Image
from docx import Document
from lxml import etree

from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import SystemMessage, HumanMessage
from langchain_core.output_parsers import StrOutputParser
from typing import List

import pymysql
import pymysql.cursors
import chainlit as cl
from google.api_core import exceptions as google_exceptions

# .env 파일이 있다면 로드
#load_dotenv()

# 실제 분석에 사용할 메인 모델
MAIN_MODEL_NAME = "gemini-3-flash-preview"

# API Key 검증용 가벼운 모델 (속도 우선)
#TEST_MODEL_NAME = "gemini-3.1-flash-lite-preview"
TEST_MODEL_NAME = "gemma-4-26b-a4b-it"

async def validate_api_key(api_key: str):
    """
    입력받은 API Key가 유효한지 flash-lite 버전의 LLM으로 빠르게 확인합니다.
    """
    try:
        # 검증용 Lite 모델 생성
        test_llm = ChatGoogleGenerativeAI(model=TEST_MODEL_NAME, google_api_key=api_key, temperature=0)
        # 1토큰 정도의 매우 짧은 응답을 유도하여 연결 확인
        await test_llm.ainvoke([HumanMessage(content="Hi")])
        return True
    except Exception as e:
        print(f"Google API 키가 유효하지 않습니다: {e}")
        return False

# --- 과목별 시스템 프롬프트 정의 (이미지 포함 버전) ---
PROMPTS = {
    "БЕЛ": """
<task>Внимателна и прецизна проверка на съдържанието на изпитните въпроси по БЕЛ</task>
<specialty>Съвременна литература, класическа литература, граматика, реторика и писане</specialty>
<caution>Текстът на въпросите и вариантите за отговор трябва да се проверяват точно според оригиналния формулировъчен изказ</caution>
<parsing_status>
    1. Таблиците и специалните символи са частично обработени (parsing)
    2. Визуалните материали като чертежи и снимки са предоставени заедно с текста
    3. Проверете съответните изображения, като се позовавате на имената на файловете в текста (напр. 'image1.jpg')
</parsing_status>
<to_dos>
    1. Проверка за фактически грешки, които могат да възникнат при интерпретацията на пасажите, граматичните правила, литературно-историческите факти и др.
    2. Проверка за възможност за неразбиране или двусмислена интерпретация на значението при анализ на изреченията
    3. Проверка дали въпросът и вариантите за отговор са ясно свързани и дали няма логически скокове
    4. Проверка дали основанието за избор на даден отговор може да бъде ясно намерено в текста
    5. Проверка за печатни грешки или тромави изреченски структури
    6. При въпроси с избор на отговор, проверка дали сред изброените варианти съществува верен отговор
    7. При въпроси с множество възможности, проверка дали съществува верен отговор чрез комбиниране на А, Б, В, Г
    8. Предоставяне на становище за преглед за всички въпроси, включително тези със свободен отговор и есеистичните въпроси
</to_dos>
<output_instructions>
    1. Отговорът трябва задължително да спазва долния Markdown формат.
    2. Преминете директно към същината без уводни думи или излишен чат.
</output_instructions>
<output_example>
## [Преглед по въпроси]
### Въпрос 1
- **Резултат от проверката**: [Подходящ / Нужда от корекция / Грешка]
- **Подробности**: (Напишете съдържание, свързано с всички инструкции от To Dos)

### Въпрос 2
...

## [Резюме на основните корекции]
- (Списък с части, които задължително трябва да бъдат коригирани)
</output_example>""",

    "Математика": """
<task>Внимателна и прецизна проверка на задачите за изпит по математика</task>
<specialty>Диференциално и интегрално смятане, Вероятности и статистика, Геометрия, Алгебра</specialty>
<caution>Текстът на задачите и вариантите за отговор трябва да се проверяват точно според оригиналния формулировъчен изказ</caution>
<parsing_status>
    1. Таблиците и специалните символи са частично обработени (parsing)
    2. Визуалните материали като чертежи и снимки са предоставени заедно с текста
    3. Проверете съответните изображения, като се позовавате на имената на файловете в текста (напр. 'image1.jpg')
</parsing_status>
<to_dos>
    1. Проверка за грешки в математическите концепции, формули и дефиниции, използвани в задачата
    2. Преглед за логически грешки или изчислителни грешки, които могат да възникнат в процеса на решаване (но не предоставяйте директно решение)
    3. Проверка дали условията на въпроса са ясни и достатъчни
    4. Проверка дали решението е еднозначно определено
    5. Проверка за случаи, в които терминологията или символите могат да бъдат интерпретирани двусмислено
    6. Проверка за печатни грешки или тромави изреченски структури
    7. При въпроси с избор на отговор, проверка дали сред изброените варианти съществува верен отговор
    8. При въпроси с множество възможности, проверка дали съществува верен отговор чрез комбиниране на А, Б, В, Г
    9. Предоставяне на становище за преглед за всички задачи, включително тези със свободен отговор и задачите за описание
</to_dos>
<output_instructions>
    1. Отговорът трябва задължително да спазва долния Markdown формат.
    2. Преминете директно към същината без уводни думи или излишен чат.
</output_instructions>
<output_example>
## [Преглед по въпроси]
### Въпрос 1
- **Резултат от проверката**: [Подходящ / Нужда от корекция / Грешка]
- **Подробности**: (Напишете съдържание, свързано с всички инструкции от To Dos)

### Въпрос 2
...

## [Резюме на основните корекции]
- (Списък с части, които задължително трябва да бъдат коригирани)
</output_example>""",

    "Английски език": """
<task>Внимателна и прецизна проверка на изпитните въпроси по английски език</task>
<specialty>Английска граматика, четене с разбиране, лексика, писане</specialty>
<caution>Текстът на въпросите и вариантите за отговор трябва да се проверяват точно според оригиналния формулировъчен изказ</caution>
<parsing_status>
    1. Таблиците и специалните символи са частично обработени (parsing)
    2. Визуалните материали като чертежи и снимки са предоставени заедно с текста
    3. Проверете съответните изображения, като се позовавате на имената на файловете в текста (напр. 'image1.jpg')
</parsing_status>
<to_dos>
    1. Проверка за наличие на граматични грешки (Grammatical errors)
    2. Проверка за неподходящ избор на лексика (Inappropriate vocabulary) или думи, които не съответстват на контекста
    3. Проверка за двусмислени (ambiguous) изрази, които могат да доведат до погрешно разбиране или многозначност при интерпретацията
    4. Проверка за логическа последователност между съдържанието на текста, въпроса и вариантите за отговор
    5. Проверка за печатни грешки (typo) или пунктуационни грешки
    6. При въпроси с избор на отговор, проверка дали сред изброените варианти съществува верен отговор
    7. При въпроси с множество възможности, проверка дали съществува верен отговор чрез комбиниране на А, Б, В, Г
    8. Предоставяне на становище за преглед за всички въпроси, включително тези със свободен отговор и задачите за есе
</to_dos>
<output_instructions>
    1. Отговорът трябва задължително да спазва долния Markdown формат.
    2. Преминете директно към същината без уводни думи или излишен чат.
</output_instructions>
<output_example>
## [Преглед по въпроси]
### Въпрос 1
- **Резултат от проверката**: [Подходящ / Нужда от корекция / Грешка]
- **Подробности**: (Напишете съдържание, свързано с всички инструкции от To Dos)

### Въпрос 2
...

## [Резюме на основните корекции]
- (Списък с части, които задължително трябва да бъдат коригирани)
</output_example>""",

    "Обществени науки": """
<task>Внимателна и прецизна проверка на изпитни въпроси по обществени науки, етика и история</task>
<specialty>Социология и култура, икономика, етика, философия, българска история, световна история, политика и право</specialty>
<caution>Текстът на въпросите и вариантите за отговор трябва да се проверяват точно според оригиналния формулировъчен изказ</caution>
<parsing_status>
    1. Таблиците и специалните символи са частично обработени (parsing)
    2. Визуалните материали като чертежи и снимки са предоставени заедно с текста
    3. Проверете съответните изображения, като се позовавате на имената на файловете в текста (напр. 'image1.jpg')
</parsing_status>
<to_dos>
    1. Проверка за фактически грешки (factual errors) в съдържанието, като правни клаузи, социално-научни концепции, етика, философия и исторически факти
    2. Проверка за наличие на изрази, които са пристрастни към определена гледна точка или са силно дискусионни
    3. Проверка за възможност за неразбиране или двусмислена интерпретация на значението при анализ на изреченията
    4. Проверка за печатни грешки в годините, имената на личностите, събитията и др.
    5. Проверка за печатни грешки (typo) или пунктуационни грешки
    6. При въпроси с избор на отговор, проверка дали сред изброените варианти съществува верен отговор
    7. При въпроси с множество възможности, проверка дали съществува верен отговор чрез комбиниране на А, Б, В, Г
    8. Предоставяне на становище за преглед за всички въпроси, включително тези със свободен отговор и тези за описание
</to_dos>
<output_instructions>
    1. Отговорът трябва задължително да спазва долния Markdown формат.
    2. Преминете директно към същината без уводни думи или излишен чат.
</output_instructions>
<output_example>
## [Преглед по въпроси]
### Въпрос 1
- **Резултат от проверката**: [Подходящ / Нужда от корекция / Грешка]
- **Подробности**: (Напишете съдържание, свързано с всички инструкции от To Dos)

### Въпрос 2
...

## [Резюме на основните корекции]
- (Списък с части, които задължително трябва да бъдат коригирани)
</output_example>""",

    "Природни науки": """
<task>Внимателна и прецизна проверка на изпитни въпроси по природни науки</task>
<specialty>Физика, химия, биология, науки за Земята</specialty>
<caution>Текстът на въпросите и вариантите за отговор трябва да се проверяват точно според оригиналния формулировъчен изказ</caution>
<parsing_status>
    1. Таблиците и специалните символи са частично обработени (parsing)
    2. Визуалните материали като чертежи и снимки са предоставени заедно с текста
    3. Проверете съответните изображения, като се позовавате на имената на файловете в текста (напр. 'image1.jpg')
</parsing_status>
<to_dos>
    1. Проверка за грешки в научните концепции, формули и дефиниции, използвани в задачата
    2. Проверка за наличие на научни грешки, заложени в смисъла на изреченията
    3. Проверка за наличие на лингвистични грешки, заложени в самото изречение
    4. Проверка за случаи, в които научните термини или изречения могат да бъдат интерпретирани двусмислено
    5. Проверка за печатни грешки (typo) или пунктуационни грешки
    6. При въпроси с избор на отговор, проверка дали сред изброените варианти съществува верен отговор
    7. При въпроси con множество възможности, проверка дали съществува верен отговор чрез комбиниране на А, Б, В, Г
    8. Предоставяне на становище за преглед за всички въпроси, включително тези със свободен отговор и тези за описание
</to_dos>
<output_instructions>
    1. Отговорът трябва задължително да спазва долния Markdown формат.
    2. Преминете директно към същината без уводни думи или излишен чат.
</output_instructions>
<output_example>
## [Преглед по въпроси]
### Въпрос 1
- **Резултат от проверката**: [Подходящ / Нужда от корекция / Грешка]
- **Подробности**: (Напишете съдържание, свързано с всички инструкции от To Dos)
- **Наличие на верен отговор**: [Няма проблем / Нужда от потвърждение]

### Въпрос 2
...

## [Резюме на основните корекции]
- (Списък с части, които задължително трябва да бъдат коригирани)
</output_example>""",
}

# --- DB 처리 함수 ---

# --- 수정된 DB 처리 함수 (과목 컬럼 분리 버전) ---

def log_usage_history(session_id, subject, is_success, error_message=None):
    """
    사용 이력을 저장합니다.
    - exam_content: 보안상 "CONTENT_NOT_STORED" 저장
    - llm_review: 성공/실패 상태 메시지 저장
    - subject: 별도 컬럼에 과목명 저장
    """
    connection = None
    try:
        connection = pymysql.connect(
            host=os.getenv("MARIADB_HOST"),
            user=os.getenv("MARIADB_USER"),
            password=os.getenv("MARIADB_PASSWORD"),
            db=os.getenv("MARIADB_DATABASE"),
            charset='utf8mb4',
            cursorclass=pymysql.cursors.DictCursor,
            port=int(os.getenv("MARIADB_PORT"))
        )
        with connection.cursor() as cursor:
            # 보안 처리
            safe_content_log = "CONTENT_NOT_STORED (Security Policy)"
            
            # 상태 메시지 (과목명은 별도 컬럼에 저장하므로 여기선 상태만 기록해도 됨)
            if is_success:
                status_log = "[SUCCESS] Review Completed"
            else:
                status_log = f"[ERROR] {error_message}"

            # 쿼리 수정: subject 컬럼 추가
            sql = """
                INSERT INTO paper_review_logs (session_id, subject, exam_content, llm_review) 
                VALUES (%s, %s, %s, %s)
            """
            cursor.execute(sql, (session_id, subject, safe_content_log, status_log))
            
        connection.commit()
        print(f"이력 저장 완료: session_id={session_id}, subject={subject}, success={is_success}")
    except pymysql.Error as e:
        print(f"이력 저장 실패: {e}")
    finally:
        if connection:
            connection.close()

# --- 파일 및 이미지 처리 함수 ---

def debug_math_structure(oMath_element):
    """수식의 XML 구조를 출력 (디버깅용)"""
    print("\n    [수식 XML 구조]:")
    xml_str = etree.tostring(oMath_element, encoding='unicode', pretty_print=True)
    # 처음 500자만 출력
    print(xml_str[:500])
    print("    ...")


def omml_to_text(oMath_element, debug=False):
    """OMML 수식을 텍스트로 변환"""
    M_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
    
    if debug:
        debug_math_structure(oMath_element)
    
    # 방법 1: m:t 요소에서 텍스트 추출
    text_parts = []
    for t in oMath_element.iter(f'{M_NS}t'):
        if t.text:
            text_parts.append(t.text)
            if debug:
                print(f"      m:t 발견: '{t.text}'")
    
    if text_parts:
        result = ''.join(text_parts)
        if debug:
            print(f"      결과 (m:t): '{result}'")
        return result
    
    # 방법 2: 모든 텍스트 노드 추출
    all_text = ''.join(oMath_element.itertext()).strip()
    if all_text:
        if debug:
            print(f"      결과 (itertext): '{all_text}'")
        return all_text
    
    if debug:
        print("      결과: (빈 문자열)")
    
    return ''

# ============================================================
# 1. OMML(수식)을 LaTeX로 변환하는 함수 정의 (먼저 실행되어야 함)
# ============================================================
def omml_to_latex(element):
    """
    OMML(Office Math Markup Language) 요소를 LaTeX 문자열로 변환하는 재귀 함수
    """
    namespaces = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    
    tag = element.tag
    local_tag = tag.split('}')[-1] # 네임스페이스 제거한 태그명

    # 1. 텍스트 (m:t)
    if local_tag == 't':
        return element.text if element.text else ""

    # 2. 분수 (m:f) -> \frac{분자}{분모}
    elif local_tag == 'f':
        num_node = element.find('m:num', namespaces)
        den_node = element.find('m:den', namespaces)
        num = omml_to_latex(num_node) if num_node is not None else ""
        den = omml_to_latex(den_node) if den_node is not None else ""
        return f"\\frac{{{num}}}{{{den}}}"

    # 3. 위첨자 (m:sSup) -> ^{...}
    elif local_tag == 'sSup':
        e_node = element.find('m:e', namespaces)
        sup_node = element.find('m:sup', namespaces)
        base = omml_to_latex(e_node) if e_node is not None else ""
        sup = omml_to_latex(sup_node) if sup_node is not None else ""
        return f"{{{base}}}^{{{sup}}}"

    # 4. 아래첨자 (m:sSub) -> _{...}
    elif local_tag == 'sSub':
        e_node = element.find('m:e', namespaces)
        sub_node = element.find('m:sub', namespaces)
        base = omml_to_latex(e_node) if e_node is not None else ""
        sub = omml_to_latex(sub_node) if sub_node is not None else ""
        return f"{{{base}}}_{{{sub}}}"

    # 5. 위아래첨자 (m:sSubSup) -> _{...}^{...}
    elif local_tag == 'sSubSup':
        e_node = element.find('m:e', namespaces)
        sub_node = element.find('m:sub', namespaces)
        sup_node = element.find('m:sup', namespaces)
        base = omml_to_latex(e_node) if e_node is not None else ""
        sub = omml_to_latex(sub_node) if sub_node is not None else ""
        sup = omml_to_latex(sup_node) if sup_node is not None else ""
        return f"{{{base}}}_{{{sub}}}^{{{sup}}}"

    # 6. 근호/루트 (m:rad) -> \sqrt[...]
    elif local_tag == 'rad':
        deg_node = element.find('m:deg', namespaces) # 제곱근의 차수
        e_node = element.find('m:e', namespaces)
        base = omml_to_latex(e_node) if e_node is not None else ""
        
        if deg_node is not None and len(deg_node) > 0:
            deg = omml_to_latex(deg_node)
            if deg.strip():
                return f"\\sqrt[{deg}]{{{base}}}"
        return f"\\sqrt{{{base}}}"
    
    # 7. 괄호/구분자 (m:d) -> \left( ... \right)
    elif local_tag == 'd':
        dPr = element.find('m:dPr', namespaces)
        begChr = '('
        endChr = ')'
        if dPr is not None:
            beg_attr = dPr.find('m:begChr', namespaces)
            end_attr = dPr.find('m:endChr', namespaces)
            if beg_attr is not None: begChr = beg_attr.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val') or '('
            if end_attr is not None: endChr = end_attr.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val') or ')'
            
        e_nodes = element.findall('m:e', namespaces)
        content = "".join([omml_to_latex(child) for child in e_nodes])
        return f"\\left{begChr} {content} \\right{endChr}"

    # 기본 재귀 호출
    else:
        text = ""
        for child in element:
            text += omml_to_latex(child)
        return text

# ============================================================
# 2. 단락 내용 추출 함수 (omml_to_latex를 호출함)
# ============================================================
def extract_paragraph_content(para_element, debug=False):
    """단락에서 텍스트와 수식을 순서대로 추출"""
    
    # 네임스페이스 정의
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    }
    
    if debug:
        print("\n  [단락 처리 시작]")
    
    para_parts = []
    
    # 단락의 모든 자식 요소 순회
    for child in para_element:
        # 수식(OMML) 발견 시
        if child.tag.endswith('oMath') or child.tag.endswith('oMathPara'):
            # 여기서 위에서 정의한 omml_to_latex 함수를 호출합니다
            latex_eq = omml_to_latex(child)
            para_parts.append(f" ${latex_eq}$ ") 
        
        # 일반 텍스트 런(w:r) 처리
        elif child.tag.endswith('r'):
            t_vals = child.findall('.//w:t', namespaces)
            for t in t_vals:
                if t.text:
                    para_parts.append(t.text)
    
    result = ''.join(para_parts).strip()
    
    if debug:
        print(f"  [단락 결과]: '{result[:100]}{'...' if len(result) > 100 else ''}'")
    
    return result


def extract_table_from_element(table_element, debug=False):
    """XML 테이블 요소에서 마크다운 표 생성"""
    W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    if debug:
        print("\n[표 처리 시작]")
    
    rows = []
    for row_idx, tr in enumerate(table_element.iter(f'{W_NS}tr')):
        row = []
        for col_idx, tc in enumerate(tr.iter(f'{W_NS}tc')):
            cell_parts = []
            
            # 셀 내의 각 단락 처리
            for para in tc.findall(f'{W_NS}p'):
                if debug:
                    print(f"  표 셀 [{row_idx},{col_idx}]:")
                para_text = extract_paragraph_content(para, debug=debug)
                if para_text:
                    cell_parts.append(para_text)
            
            cell_text = ' '.join(cell_parts).strip()
            row.append(cell_text if cell_text else ' ')
        
        if row:
            rows.append(row)
    
    if not rows:
        return None
    
    # 마크다운 표 생성
    markdown_table = []
    
    if rows:
        header = ' | '.join(rows[0])
        markdown_table.append(f"| {header} |")
        
        separator = ' | '.join(['---'] * len(rows[0]))
        markdown_table.append(f"| {separator} |")
        
        for row in rows[1:]:
            while len(row) < len(rows[0]):
                row.append(' ')
            row_text = ' | '.join(row[:len(rows[0])])
            markdown_table.append(f"| {row_text} |")
    
    if debug:
        print("[표 처리 완료]")
    
    return '\n'.join(markdown_table)


def extract_textbox_content(textbox_element, debug=False):
    """글상자 요소에서 텍스트, 표, 수식을 모두 추출"""
    W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    content_parts = []
    
    if debug:
        print("\n" + "="*60)
        print("글상자 내용 추출 시작")
        print("="*60)
    
    # w:txbxContent 찾기
    for txbxContent in textbox_element.iter(f'{W_NS}txbxContent'):
        for child in txbxContent:
            # 표 발견
            if child.tag == f'{W_NS}tbl':
                table_md = extract_table_from_element(child, debug=debug)
                if table_md:
                    content_parts.append('\n' + table_md + '\n')
            
            # 단락 발견
            elif child.tag == f'{W_NS}p':
                para_text = extract_paragraph_content(child, debug=False)
                if para_text:
                    content_parts.append(para_text)
    
    result = '\n\n'.join(content_parts)
    
    if debug:
        print("="*60)
        print(f"글상자 내용 추출 완료 ({len(result)}자)")
        print("="*60 + "\n")
    
    return result


def convert_docx_to_text_with_images(file_path, media_dir='./extracted_media', debug=False):
    """글상자를 포함한 DOCX를 마크다운으로 변환"""
    
    if not os.path.exists(file_path):
        return "파일을 찾을 수 없습니다."
    
    try:
        print(f"변환 시작: {file_path}")
        
        doc = Document(file_path)
        textbox_map = {}
        all_textboxes = []
        
        temp_path = file_path.replace('.docx', '_temp.docx')
        
        # 네임스페이스 정의
        VML_NS = '{urn:schemas-microsoft-com:vml}'
        WPS_NS = '{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}'
        
        textbox_count = 0
        for para in doc.paragraphs:
            para_textboxes = []
            para_element = para._element
            
            # VML textbox
            for textbox in para_element.iter(f'{VML_NS}textbox'):
                textbox_count += 1
                if debug:
                    print(f"\n>>> 글상자 #{textbox_count} (VML) 처리 중...")
                
                content = extract_textbox_content(textbox, debug=debug)
                
                if content.strip():
                    para_textboxes.append(content)
                    all_textboxes.append(content)
            
            # wps:txbx
            for txbx in para_element.iter(f'{WPS_NS}txbx'):
                textbox_count += 1
                if debug:
                    print(f"\n>>> 글상자 #{textbox_count} (WPS) 처리 중...")
                
                content = extract_textbox_content(txbx, debug=debug)
                
                if content.strip():
                    para_textboxes.append(content)
                    all_textboxes.append(content)
            
            if para_textboxes:
                textbox_map[id(para)] = para_textboxes
                for content in para_textboxes:
                    preview = content[:80].replace('\n', ' ') + "..." if len(content) > 80 else content.replace('\n', ' ')
                    if not debug:  # debug 모드가 아닐 때만 간단한 요약 출력
                        print(f"  ✓ 글상자 발견: {preview}")
        
        print(f"\n총 {len(all_textboxes)}개의 글상자를 발견했습니다.")
        
        if not all_textboxes:
            print("글상자를 찾지 못했습니다. Pandoc만으로 변환합니다.")
            markdown = pypandoc.convert_file(
                file_path,
                to='markdown',
                format='docx',
                extra_args=[
                    '--wrap=none',
                    '--standalone',
                    f'--extract-media={media_dir}'
                ]
            )
            return markdown
        
        # 마커 삽입
        marker_to_content = {}
        inserted_count = 0
        
        for para in doc.paragraphs:
            para_id = id(para)
            if para_id in textbox_map:
                textboxes = textbox_map[para_id]
                
                for content in textboxes:
                    marker = f"@@TEXTBOX{inserted_count + 1}@@"
                    marker_to_content[marker] = content
                    
                    if para.runs:
                        para.runs[-1].text += f" {marker} "
                    else:
                        para.add_run(f" {marker} ")
                    
                    inserted_count += 1
                    if not debug:
                        preview = content[:50].replace('\n', ' ') + "..." if len(content) > 50 else content.replace('\n', ' ')
                        print(f"  ✓ {marker} 삽입됨: {preview}")
        
        print(f"\n마커 삽입 완료: {inserted_count}개")
        
        doc.save(temp_path)
        
        # Pandoc 변환
        markdown = pypandoc.convert_file(
            temp_path,
            to='markdown',
            format='docx',
            extra_args=[
                '--wrap=none',
                '--standalone',
                f'--extract-media={media_dir}'
            ]
        )
        
        # 마커 확인
        found = sum(1 for m in marker_to_content if m in markdown)
        print(f"  발견된 마커: {found}/{len(marker_to_content)}개")
        
        # 마커 치환
        replaced_count = 0
        
        for marker, content in marker_to_content.items():
            if marker in markdown:
                if '|' in content and '---' in content:
                    formatted_content = f"\n\n> 📦 **[글상자]**\n\n{content}\n\n"
                else:
                    lines = content.split('\n')
                    formatted_lines = ['> ' + line if line.strip() else '>' for line in lines]
                    formatted_content = f"\n\n> 📦 **[글상자]**\n" + '\n'.join(formatted_lines) + "\n\n"
                
                markdown = markdown.replace(marker, formatted_content)
                replaced_count += 1
                if not debug:
                    print(f"  ✓ {marker} 치환 완료")
        
        # 미치환 글상자 처리
        if replaced_count < len(marker_to_content):
            print(f"\n⚠ {len(marker_to_content) - replaced_count}개 글상자를 문서 끝에 추가")
            markdown += "\n\n---\n\n## 📦 글상자 내용 (원래 위치 확인 불가)\n\n"
            
            for i, (marker, content) in enumerate(marker_to_content.items(), 1):
                if marker not in markdown:
                    if '|' in content and '---' in content:
                        markdown += f"\n### 글상자 {i}\n\n{content}\n\n"
                    else:
                        lines = content.split('\n')
                        formatted_lines = ['> ' + line if line.strip() else '>' for line in lines]
                        markdown += f"\n### 글상자 {i}\n\n" + '\n'.join(formatted_lines) + "\n\n"
        
        if os.path.exists(temp_path):
            os.remove(temp_path)
        
        print(f"\n✅ 변환 완료: {replaced_count}/{len(marker_to_content)}개 글상자 삽입됨")
        
        return markdown
        
    except Exception as e:
        temp_path = file_path.replace('.docx', '_temp.docx')
        if os.path.exists(temp_path):
            os.remove(temp_path)
        import traceback
        return f"오류 발생: {e}\n\n{traceback.format_exc()}"

def preprocess_content(text):
    """
    Pandoc으로 변환된 텍스트의 노이즈를 제거하고 이미지 태그를 정규화합니다.
    1. 불필요한 ASCII 테이블 선(+---+) 제거
    2. 복잡한 이미지 경로 및 속성 제거 -> [이미지: image1.jpg] 형태로 통일
    3. 과도한 공백 제거
    """

    # 1. ASCII 테이블의 구분선 (+----+) 제거
    # 줄의 시작이 +이고, 그 뒤로 -나 +로만 이루어진 줄을 찾아 제거
    text = re.sub(r'^\+[-+]+$', '', text, flags=re.MULTILINE)

    # 2. 이미지 태그 정규화 및 확장자 강제 변경 (.jpg)
    # 패턴: ![설명](경로){속성} 또는 ![설명](경로)
    def clean_image_tag(match):
        # 정규식 그룹 1번: 이미지 경로 (예: ./extracted_media/media/image1.png)
        full_path = match.group(1)
        
        # 파일명만 추출 (image1.png)
        filename = os.path.basename(full_path)
        name, ext = os.path.splitext(filename)
        
        # 사용자의 로직에 맞춰 모든 확장자를 .jpg로 변경하여 표기
        new_filename = f"{name}.jpg"
        
        # LLM이 인식하기 쉬운 형태로 변경
        return f"\n[참고 이미지: {new_filename}]\n"

    # Markdown 이미지 링크 패턴 매칭 및 치환
    # ![](path){opt} 또는 ![](path) 형태 모두 대응
    text = re.sub(r'!\[.*?\]\((.*?)\)(?:\{.*?\})?', clean_image_tag, text)

    # 3. 테이블 파이프(|)로 인한 불필요한 공백 및 줄바꿈 정리
    # 연속된 3개 이상의 줄바꿈을 2개로 축소
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text

def resize_image_if_needed(img, max_width=640):
    """이미지 가로 크기가 max_width보다 크면 비율 유지하며 리사이징"""
    width, height = img.size
    if width > max_width:
        ratio = max_width / width
        new_height = int(height * ratio)
        img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
    return img

def convert_to_grayscale(img):
    """이미지를 그레이스케일로 변환"""
    if img.mode != 'L':
        img = img.convert('L')
    return img

def convert_images_to_jpg(directory, max_width=640):
    """EMF, PNG, BMP, WMF 파일을 JPG로 변환 (리사이징 및 그레이스케일)"""
    directory = Path(directory)
    if not directory.exists(): return
    extensions = ['.emf', '.png', '.bmp', '.wmf']
    for ext in extensions:
        for file_path in directory.glob(f'*{ext}'):
            try:
                with Image.open(file_path) as img:
                    if img.mode in ('RGBA', 'LA', 'P'):
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                        img = background
                    elif img.mode not in ('RGB', 'L'):
                        img = img.convert('RGB')
                    
                    img = resize_image_if_needed(img, max_width)
                    img = convert_to_grayscale(img)
                    
                    new_file_path = file_path.with_suffix('.jpg')
                    img.save(new_file_path, 'JPEG', quality=75, optimize=True)
            except Exception as e:
                print(f"이미지 변환 실패: {file_path.name} - {str(e)}")

def resize_existing_jpg(directory, max_width=640):
    """기존 JPG/JPEG 파일을 리사이징 및 그레이스케일 변환"""
    directory = Path(directory)
    if not directory.exists(): return
    jpg_files = list(directory.glob('*.jpg')) + list(directory.glob('*.jpeg'))
    for file_path in jpg_files:
        try:
            with Image.open(file_path) as img:
                needs_processing = img.size[0] > max_width or img.mode != 'L'
                if needs_processing:
                    img = resize_image_if_needed(img, max_width)
                    img = convert_to_grayscale(img)
                    img.save(file_path, 'JPEG', quality=75, optimize=True)
        except Exception as e:
            print(f"JPG 처리 실패: {file_path.name} - {str(e)}")

def encode_images_to_base64(directory):
    """디렉토리의 모든 JPG/JPEG 파일을 Base64로 인코딩 (파일명 포함)"""
    directory = Path(directory)
    if not directory.exists(): return []
    encoded_data = []
    jpg_files = sorted(list(directory.glob('*.jpg'))) + sorted(list(directory.glob('*.jpeg')))
    for file_path in jpg_files:
        try:
            with open(file_path, 'rb') as image_file:
                encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
                encoded_data.append({
                    "filename": file_path.name,  # ← 파일명 추가
                    "b64": encoded_string
                })
        except Exception as e:
            print(f"Base64 인코딩 실패: {file_path.name} - {str(e)}")
    return encoded_data

# --- Chainlit 앱 로직 ---

@cl.on_chat_start
async def on_chat_start():
    res = await cl.AskUserMessage(
        content="Здравейте! Аз съм Вашият асистент за преглед на училищни изпитни въпроси.\n"
                "Системата е проектирана така, че нито съдържанието на въпросите, нито резултатите от проверката да се съхраняват на сървъра, така че можете да я използвате без никакви притеснения.\n\n"
                "Моля, въведете Вашия **Google Gemini API ключ**.\n\n"
                "API ключът също не се записва на сървъра. Въведете го уверено.\n"
                "С един безплатен API ключ можете да качвате и проверявате цели файлове с изпитни въпроси до **20 пъти на ден**.\n"
                "Можете да получите безплатен Google Gemini API ключ на адрес: https://aistudio.google.com/app/api-keys", 
        timeout=600
    ).send()

    if res:
        user_api_key = res["output"].strip()
        
        # 검증 중 메시지
        msg = cl.Message(content=f"**{TEST_MODEL_NAME}** Проверка на връзката с модела чрез API ключа...", author="문제 검토 도우미")
        await msg.send()
        
        # Lite 모델로 빠르게 검증
        is_valid = await validate_api_key(user_api_key)
        
        # 수정안: 재입력 루프 없이 검증 실패 시 안내만 하고 종료
        if is_valid:
            cl.user_session.set("user_api_key", user_api_key)
            msg.content = "✅ API ключът е потвърден."
            await msg.update()
        else:
            msg.content = "❌ Невалиден API ключ. Моля, презаредете страницата и опитайте отново."
            await msg.update()
            return
    else:
        # 타임아웃
        await cl.Message(content="Времето за въвеждане изтече. Моля, презаредете страницата.").send()
        return

    # 2. 인사말 및 영역 선택
    await cl.Message(content=f"Модел за анализ: {MAIN_MODEL_NAME}", author="문제 검토 도우미").send()

    actions = [
        cl.Action(name="subject_select", value=subject, label=subject, payload={"subject": subject})
        for subject in PROMPTS.keys()
    ]
    await cl.Message(
        content="Моля, изберете предмет за преглед.",
        actions=actions,
        author="문제 검토 도우미"
    ).send()


@cl.action_callback("subject_select")
async def on_subject_select(action: cl.Action):
    """과목 선택 시 세션에 저장하고 안내를 표시합니다."""
    subject = action.payload["subject"]
    cl.user_session.set("subject", subject)
    cl.user_session.set("session_id", cl.user_session.get("id"))
    await cl.Message(content=f"Избрахте предмет **{subject}**.").send()
    await cl.Message(content="Ще ви запозная с начина на употреба.\n"
            "1. Изберете 'Запис като' и запишете файла във формат **Word документ (*.docx)**.\n"
            "2. **Отворете преобразувания файл (.docx), изтрийте излишно генерираните номера на задачи и го запазете.** Това е необходимо, тъй като преобразуването от HWP към Word не винаги е перфектно.\n"
            "3. Моля, качете преобразувания файл с изпитни въпроси (.docx).", author="문제 검토 도우미").send()
    await action.remove()


@cl.on_message
async def on_message(message: cl.Message):
    # on_message 함수 내부에서
    user_api_key = cl.user_session.get("user_api_key")
    if not user_api_key:
        await cl.Message(content="API ключът липсва. Моля, презаредете страницата.").send()
        return

    llm = ChatGoogleGenerativeAI(
        model=MAIN_MODEL_NAME,
        google_api_key=user_api_key,
        thinking_level="high"
    )

    subject = cl.user_session.get("subject")
    if not subject:
        await cl.Message(content="Моля, изберете предмет за преглед.", author="문제 검토 도우미").send()
        return

    files = [f for f in message.elements if isinstance(f, cl.File)]
    if not files:
        await cl.Message(content="Моля, качете преобразувания файл с изпитни въпроси (.docx).", author="문제 검토 도우미").send()
        return

    uploaded_file = files[0]
    if not uploaded_file.name.lower().endswith(".docx"):
        await cl.Message(content="Моля, качете файл във формат .docx.", author="문제 검토 도우미").send()
        return

    session_id = cl.user_session.get("session_id")
    media_dir = f"./extracted_media_{session_id}"
    
    if os.path.exists(media_dir):
        shutil.rmtree(media_dir)
    os.makedirs(media_dir)

    try:
        await cl.Message(content=f"'{uploaded_file.name}' файлът с изпитни въпроси ще бъде прегледан.\n\nТекстът и изображенията ще бъдат обработени заедно. Това, че LLM има определени ограничения при обработка на изображения, може да доведе до непълно извличане на информация от тях. Моля, прочетете коментарите за преглед.", author="문제 검토 도우미").send()
        
        extracted_text = await cl.make_async(convert_docx_to_text_with_images)(uploaded_file.path, media_dir, debug=False)
        
        if extracted_text.startswith("Грешка:"):
            # 변환 실패 로그 저장
            await cl.make_async(log_usage_history)(session_id, subject, False, error_message=f"File Convert Error: {extracted_text}")
            await cl.Message(content=f"Възникна проблем при обработка на файла.\n\n{extracted_text}", author="문제 검토 도우미").send()
            return
        
        extracted_text = preprocess_content(extracted_text)
        image_processing_dir = os.path.join(media_dir, 'media')

        await cl.make_async(convert_images_to_jpg)(image_processing_dir)
        await cl.make_async(resize_existing_jpg)(image_processing_dir)
        encoded_images = await cl.make_async(encode_images_to_base64)(image_processing_dir)

        system_role = PROMPTS[subject]
        
        human_contents = [{"type": "text", "text": f"시험지 본문 내용: {extracted_text}"}]

        for item in encoded_images:
            # 어떤 이미지인지 LLM에게 명확히 알려줌
            human_contents.append({
                "type": "text",
                "text": f"[이미지 파일명: {item['filename']}]"
            })
            human_contents.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{item['b64']}"}
            })
        
        messages = [
            SystemMessage(content=system_role),
            HumanMessage(content=human_contents)
        ]

        msg = cl.Message(content="", author="검토 의견")
        await msg.send()

        full_response = []
        async with cl.Step(name="Преглед на изпитните въпроси..."):
            async for chunk in llm.astream(messages):
                if isinstance(chunk.content, list):
                    # 리스트인 경우 텍스트만 추출
                    content = ''.join([
                        item.get('text', '') if isinstance(item, dict) else str(item)
                        for item in chunk.content
                    ])
                elif isinstance(chunk.content, str):
                    content = chunk.content
                else:
                    # 예상치 못한 타입은 문자열로 변환
                    content = str(chunk.content)
                
                if content:  # 빈 문자열이 아닐 때만 처리
                    await msg.stream_token(content)
                    full_response.append(content)
            
            await msg.update()

        # [수정됨] 정상 완료 시 로그 저장 (성공=True, 에러메시지=None)
        await cl.make_async(log_usage_history)(session_id, subject, True)

    except google_exceptions.InvalidArgument as e:
        error_msg = str(e)
        if "request is too large" in error_msg or "token" in error_msg.lower():
            user_msg = "Грешка: Въведеното съдържание е твърде дълго. Надвишен е максималният лимит от токени, които API може да обработи. Моля, разделете DOCX файла на няколко части и ги качете поотделно."
            await cl.Message(content=user_msg, author="오류").send()
        else:
            await cl.Message(content=f"Възникна грешка при обработката на API: {e}", author="오류").send()
        
        # [수정됨] API 오류 시 로그 저장 (성공=False, 에러메시지 기록)
        await cl.make_async(log_usage_history)(session_id, subject, False, error_message=f"Google API Error: {error_msg}")

    except Exception as e:
        error_msg = str(e)
        await cl.Message(content=f"ъзникна неочаквана грешка по време на обработката: {e}", author="문제 검토 도우미").send()
        
        # [수정됨] 일반 오류 시 로그 저장 (성공=False, 에러메시지 기록)
        await cl.make_async(log_usage_history)(session_id, subject, False, error_message=f"General Exception: {error_msg}")

    finally:
        if os.path.exists(media_dir):
            shutil.rmtree(media_dir)
            print(f"임시 디렉토리 삭제 완료: {media_dir}")
